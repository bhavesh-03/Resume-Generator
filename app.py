from flask import Flask, request, render_template, send_file
import joblib
import pandas as pd
from transformers import BertTokenizer, BertModel
from docx import Document

app = Flask(__name__)

# Load the model and tokenizer
classifier = joblib.load('model/random_forest_model.pkl')
tokenizer = joblib.load('model/bert_tokenizer.pkl')
bert_model = joblib.load('model/bert_model.pkl')

def get_bert_embeddings(text, tokenizer, model):
    inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=128)
    outputs = model(**inputs)
    return outputs.pooler_output.squeeze().detach().numpy()

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume(name, job_title, responsibilities, qualifications, experience, location, company, coding_profile_link):
    # Generate skills based on the provided job title
    job_desc = f"{job_title} {responsibilities} {qualifications} {experience}"
    embedding = get_bert_embeddings(job_desc, tokenizer, bert_model)
    skills = classifier.predict([embedding])[0]
    
    # Initialize a new Word document
    doc = Document()

    # Define styles
    styles = doc.styles
    
    # Heading style
    heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(18)
    heading_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    heading_style.paragraph_format.space_after = Pt(12)
    
    # Subheading style
    subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.name = 'Arial'
    subheading_style.font.size = Pt(14)
    subheading_style.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
    subheading_style.font.bold = True
    subheading_style.paragraph_format.space_before = Pt(12)
    subheading_style.paragraph_format.space_after = Pt(6)
    
    # Normal text style
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Add sections to the document
    doc.add_paragraph(name, style='CustomHeading').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(job_title, style='CustomSubheading').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Professional Summary', style='CustomSubheading')
    doc.add_paragraph(f"{responsibilities}", style='CustomNormal')
    
    doc.add_paragraph('Qualifications', style='CustomSubheading')
    doc.add_paragraph(f"{qualifications}", style='CustomNormal')
    
    doc.add_paragraph('Experience', style='CustomSubheading')
    doc.add_paragraph(f"{experience}", style='CustomNormal')
    
    doc.add_paragraph('Contact Information', style='CustomSubheading')
    contact_info = doc.add_paragraph(style='CustomNormal')
    contact_info.add_run(f"Location: {location}\n")
    contact_info.add_run(f"Company: {company}\n")
    contact_info.add_run(f"Coding Profile: {coding_profile_link}")
    
    # Add skills to the document
    doc.add_paragraph('Skills', style='CustomSubheading')
    skills_list = skills.split('•')  # Split by bullet points
    for skill in skills_list:
        skill = skill.strip()  # Remove leading/trailing whitespace
        if skill:  # Check if the skill is not empty
            skill_items = skill.split(',')  # Split by commas
            for item in skill_items:
                item = item.strip()  # Remove leading/trailing whitespace
                if item:  # Check if the item is not empty
                    skill_para = doc.add_paragraph(style='CustomNormal')
                    skill_para.add_run('• ').bold = True
                    skill_para.add_run(item)
                    skill_para.paragraph_format.left_indent = Pt(18)
                    skill_para.paragraph_format.space_after = Pt(3)

    # Save the document
    file_name = f"{name}_{job_title}_resume.docx"
    doc.save(file_name)
    return file_name

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/generate_resume', methods=['POST'])
def generate_resume_route():
    name = request.form['name']
    job_title = request.form['job_title']
    responsibilities = request.form['responsibilities']
    qualifications = request.form['qualifications']
    experience = request.form['experience']
    location = request.form['location']
    company = request.form['company']
    coding_profile_link = request.form['coding_profile_link']

    resume_file = generate_resume(name, job_title, responsibilities, qualifications, experience, location, company, coding_profile_link)
    return send_file(resume_file, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
